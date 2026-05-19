"""
Build the comprehensive project documentation as a Word .docx.
Output: outputs/Capstone8_Full_Documentation.docx (+ Desktop copy)
"""
import os, shutil
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x00, 0x30, 0x87)
BLUE = RGBColor(0x02, 0x77, 0xBD)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def H1(text):
    p = doc.add_heading(level=1)
    r = p.add_run(text); r.font.color.rgb = NAVY; r.font.size = Pt(18); r.bold = True
    return p

def H2(text):
    p = doc.add_heading(level=2)
    r = p.add_run(text); r.font.color.rgb = BLUE; r.font.size = Pt(14); r.bold = True
    return p

def H3(text):
    p = doc.add_heading(level=3)
    r = p.add_run(text); r.font.color.rgb = NAVY; r.font.size = Pt(12); r.bold = True
    return p

def para(text, italic=False, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text); r.italic = italic; r.bold = bold; r.font.size = Pt(size)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix + ": "); r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def make_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True; run.font.size = Pt(10); run.font.color.rgb = RGBColor(255, 255, 255)
        shade_cell(hdr[i], "003087")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t

# ══════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Lowe's Store-Level Sales Target Model")
r.font.size = Pt(26); r.bold = True; r.font.color.rgb = NAVY

s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("A Supervised Machine Learning Approach to Replace the 'Peanut Butter Spread'")
r.font.size = Pt(13); r.italic = True; r.font.color.rgb = BLUE

doc.add_paragraph()
for line, sz in [("Comprehensive Project Documentation", 15),
                 ("Capstone 8  ·  Group 8", 13),
                 ("IIM Calcutta  ·  APAL02", 12),
                 ("Models · Features · Methodology · Results · Reflections", 11)]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = p.add_run(line); rr.font.size = Pt(sz); rr.bold = (sz >= 13)
    if sz < 12: rr.font.color.rgb = GREY

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("GitHub: github.com/nishthaspeakx/capstone_8").font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════
H1("1.  Executive Summary")
para("Lowe's currently sets each store's annual sales target with a 'peanut butter spread' — "
     "it divides the total company plan across stores using each store's 3-year historical sales "
     "share. This ignores local market dynamics (population growth, income shifts, housing age, "
     "competitive pressure), so under the current method roughly 25.83% of store-weeks miss their "
     "target by more than ±10%, forcing planners into a year of manual overrides.")
para("This project replaces that heuristic with a supervised machine-learning model. We trained "
     "and benchmarked four models on 269,412 store-week records (1,727 stores, FY2023–FY2025), "
     "engineered a disciplined, leakage-free feature set, and finalized a 16-feature XGBoost model. "
     "We then built a Store-Level Accuracy Loop that flags exactly which stores are under- or "
     "over-targeted, plus a Role-of-Store segmentation that turns one model into five planning "
     "playbooks. The deliverable is an interactive Streamlit application.")

H3("Final Model Scorecard")
make_table(
    ["Metric", "Result", "Success Criterion", "Status"],
    [["Overall WAPE", "7.05%", "< 8%", "PASS"],
     ["Bias", "-0.97%", "within ±2%", "PASS"],
     ["R²", "0.977", "> 0.88", "PASS"],
     ["Beat naive baseline", "7.05% vs 8.27%", "beat lag-52", "PASS (15% better)"],
     ["Well-calibrated stores", "1,636 of 1,727 (94.7%)", "maximize", "PASS"]],
    widths=[2.0, 2.0, 1.8, 1.6])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 2. THE MODELS — WHAT & WHY
# ══════════════════════════════════════════════════════════════
H1("2.  The Models — What We Used and Why")
para("This is a supervised regression problem: we have a clear numeric target (Actual Sales USD) "
     "and labelled history, so the model learns input→output patterns and predicts a number. "
     "Clustering (unsupervised) is used only afterwards as a strategic overlay, not as the core model.")

H2("2.1  XGBoost — Primary Model (Selected)")
para("Extreme Gradient Boosting builds 500 decision trees sequentially. Tree 1 makes predictions, "
     "computes the errors (residuals); Tree 2 is trained on those errors to reduce them; Tree 3 "
     "reduces what remains, and so on. Each tree corrects the mistakes of all previous trees — "
     "this is 'boosting'.")
para("Why we chose it:", bold=True)
bullet("Best WAPE of all candidates (7.05%) and near-zero bias (-0.97%).")
bullet("Handles non-linear interactions between market features automatically (e.g., 'high income AND new housing AND low competition' is learned without us specifying it).")
bullet("Robust to mixed feature scales — no need to normalize.")
bullet("Built-in feature importance, which is essential for explaining targets to planners.")
para("Hyperparameters used:", bold=True)
make_table(["Parameter", "Value", "Reason"],
    [["n_estimators", "500", "Enough trees for refinement; early stopping prevents overfit"],
     ["learning_rate", "0.05", "Small steps → stable, accurate convergence"],
     ["max_depth", "7", "Limits tree complexity, controls overfitting"],
     ["subsample", "0.8", "Row randomness → generalization"],
     ["colsample_bytree", "0.8", "Feature randomness → de-correlated trees"],
     ["early_stopping", "50 rounds", "Stop when validation stops improving"]],
    widths=[1.7, 1.2, 4.0])

H2("2.2  LightGBM — Secondary Benchmark")
para("Same gradient-boosting concept, but it grows trees leaf-by-leaf (always splitting the leaf "
     "with the largest loss reduction) rather than level-by-level. It is faster on large datasets. "
     "Result: WAPE 7.07%, essentially tied with XGBoost — this cross-validates that the result is "
     "model-robust, not a fluke of one algorithm.")

H2("2.3  Ridge Regression — Linear Baseline")
para("A linear regression with an L2 penalty (α·Σβ²) that shrinks coefficients to prevent "
     "instability from correlated features. It is a sanity check: if gradient boosting could not "
     "beat a regularized linear model, the added complexity would not be justified. Ridge scored "
     "WAPE 7.11% — competitive, which tells us the relationship is fairly linear once lag features "
     "are present, but XGBoost still edges it out and captures the non-linear market effects.")

H2("2.4  Naive Baseline — Lag-52")
para("The simplest possible 'model': this week's sales = the same week last year. WAPE 8.27%. "
     "Any real model must beat this to justify its existence. Our final model is 15% better, "
     "which is the headline improvement story.")

H2("2.5  The Logic Behind the Models — Mathematical Intuition")
para("This section explains WHY these models work, in plain terms, for the technical Q&A.")

H3("Supervised Regression — the core idea")
para("We are learning a function f such that f(X) ≈ y, where X is the 16-feature vector for a "
     "store-week and y is that week's Actual Sales USD. The model is given thousands of historical "
     "(X, y) pairs and adjusts itself to minimize the gap between its prediction ŷ = f(X) and the "
     "true y. 'Supervised' simply means every training example carries the correct answer (the label).")

H3("Why Gradient Boosting (XGBoost / LightGBM)")
para("A single decision tree asks a series of yes/no questions ('Is lag_1 > $900K?', then "
     "'Is housing_new_share > 5%?') and lands each store-week in a leaf with a predicted value. "
     "One tree is weak — it either oversimplifies or overfits. Gradient boosting fixes this by "
     "building an ENSEMBLE sequentially:")
bullet("Tree 1 makes a rough prediction. We measure its error for every row (the residual = actual − predicted).")
bullet("Tree 2 is trained NOT on sales, but on Tree 1's residuals — it learns to predict 'where Tree 1 was wrong'.")
bullet("Tree 3 corrects what Tree 1+2 still get wrong, and so on for 500 trees.")
bullet("Final prediction = Tree 1 + 0.05×Tree 2 + 0.05×Tree 3 + … (the 0.05 is the learning rate — small steps avoid overshooting).")
para("Formally, at each step m the model adds the tree h_m that best reduces the loss L: "
     "F_m(x) = F_{m-1}(x) + ν · h_m(x), where h_m is fit to the negative gradient of L with "
     "respect to F_{m-1} (hence 'gradient' boosting). With squared-error loss the negative "
     "gradient is exactly the residual, which is why the intuition above is precise, not just "
     "an analogy.")
para("Why it beats a linear model here: sales depend on INTERACTIONS — e.g., high income matters "
     "more in a fast-growing housing market with low competition. Trees capture these "
     "'if-and-and' combinations automatically; a linear model would need every interaction "
     "term specified by hand.")

H3("Why XGBoost over LightGBM (both gradient boosting)")
para("XGBoost grows each tree level-by-level (all nodes at a depth split before going deeper) and "
     "adds an explicit regularization term (γ, λ) penalizing tree complexity. LightGBM grows "
     "leaf-by-leaf (always split the leaf with the biggest loss gain) — faster, but slightly more "
     "prone to overfitting on smaller signals. On our holdout XGBoost edged ahead on WAPE and "
     "bias, so it is the production choice; LightGBM's near-identical score is our robustness check.")

H3("Why Ridge as the linear control")
para("Ordinary linear regression becomes unstable when inputs are correlated (our lag features "
     "are highly correlated with each other). Ridge adds an L2 penalty λ·Σβ² to the loss, which "
     "shrinks coefficients toward zero and stabilizes them. It answers the question: 'Is the "
     "expensive non-linear model actually necessary?' Ridge reaching 7.11% tells us lags make the "
     "problem nearly linear, but XGBoost's edge confirms the market interactions are real and worth modeling.")

H3("Why WAPE as the headline metric")
para("WAPE = Σ|actual − predicted| / Σ(actual). Unlike MAPE, it weights each error by sales "
     "volume, so a $1K miss at a $50M store does not count the same as a $1K miss at a $1M store "
     "— exactly how a planner thinks about dollar risk. Bias = (Σpred − Σactual)/Σactual measures "
     "systematic direction (does the model habitually over- or under-commit?). R² measures the "
     "share of sales variation explained. We require all three to pass, not just one.")

H3("Head-to-Head Comparison (Holdout Set)")
comp = pd.read_csv("outputs/model_comparison.csv", index_col=0)
rows = []
for idx, row in comp.iterrows():
    rows.append([idx, f"{row['WAPE']:.2f}%", f"{row['Bias']:+.2f}%",
                 f"${row['RMSE']:,.0f}", f"${row['MAE']:,.0f}", f"{row['R2']:.4f}"])
make_table(["Model", "WAPE", "Bias", "RMSE", "MAE", "R²"], rows,
           widths=[1.6, 0.9, 0.9, 1.3, 1.3, 0.9])
para("Holdout = rolling time split, last ~15% of dates (Sep 19, 2025 → Jan 30, 2026, all 1,727 "
     "stores, 22,439 rows). XGBoost selected as the production model.", italic=True, size=9)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 3. THE 16 FEATURES
# ══════════════════════════════════════════════════════════════
H1("3.  The Final 16 Features — What and Why")
para("We began with 35 Phase-1 candidate features (already filtered from 194 raw columns), then "
     "ran a criticality analysis and finalized 16 high-criticality features. Critically, we "
     "deliberately removed the rolling-average features (roll_4, roll_13) and the calendar "
     "features (Year, Fiscal Week). The rolling averages were so dominant (~94% combined "
     "importance) that they masked every market signal. Removing them made the model more "
     "explainable to planners AND improved bias to -0.97% (near-perfect).")

H2("3.1  Lag / Momentum Features (4)")
para("Engineered from Actual Sales USD using groupby('Store ID').shift(n). Every lag uses a "
     "minimum shift of 1 so the model can never see the week it is predicting (leakage-safe).")
make_table(["Feature", "Formula", "Logic — Why It's Used"],
    [["lag_1", "shift(1)", "Sales 1 week ago. Strongest single predictor (corr 0.982) — captures very recent store momentum."],
     ["lag_4", "shift(4)", "Sales ~1 month ago. Monthly rhythm: pay cycles, monthly promotions, billing."],
     ["lag_13", "shift(13)", "Sales ~1 quarter ago. Quarter-over-quarter retail dynamics."],
     ["lag_52", "shift(52)", "Sales same week last year. Pure annual seasonality — the most reliable seasonal anchor."]],
    widths=[1.0, 1.0, 4.8])

H2("3.2  Demand / Housing — Raw (5)")
make_table(["Feature", "Logic — Why It's Used"],
    [["CYE Total Households", "Core addressable market. More households = more potential customers."],
     ["CYE Household Density HH/SqMi", "Distinguishes dense walkable trade areas from sprawling car-based ones — different basket behavior."],
     ["CYE Median Household Income", "Best single summary of purchasing power. Median is robust to outliers."],
     ["CYE Total Housing Units", "Housing stock incl. vacant/new — captures renovation/construction demand beyond current occupants."],
     ["CYE Total Population", "Market-size complement to households (captures household-size variation)."]],
    widths=[2.2, 4.6])

H2("3.3  Demand / Housing — Engineered (3)")
para("Each replaces multiple raw columns to avoid multicollinearity (the original income brackets "
     "sum to 100%, housing decades sum to 100% — using all of them destabilizes the model).")
make_table(["Feature", "Formula", "Logic"],
    [["income_affluent", "%150–250K + %250K+", "Pro/contractor proxy — wealthy areas drive remodeling & high-value purchases."],
     ["housing_new_share", "%2000s + %2010s", "New construction → new-homeowner finishing/landscaping demand. Highest static-feature correlation."],
     ["housing_old_share", "%pre-1949 + %1950s + %1960s", "Old stock → continuous repair demand (plumbing, electrical, roofing)."]],
    widths=[1.6, 1.9, 3.3])

H2("3.4  Competition (3)")
make_table(["Feature", "Logic"],
    [["Sister Store Count (TA)", "Self-cannibalization. 3 Lowe's in one trade area share demand — corrects raw market size."],
     ["total_competitor_ta", "Engineered: sum of all 17 competitor counts. Single composite of competitive intensity."],
     ["Wallflowers Depot Count (TA)", "Most prevalent competitor (mean 4.68, only 1% zeros) — broad signal."]],
    widths=[2.2, 4.6])

H2("3.5  Store / Market (1)")
make_table(["Feature", "Logic"],
    [["Urbanicity (encoded)", "7-level density (Metropolis→Remote). Master segmentation variable — Metropolis stores avg $1.63M/wk vs Large City $930K."]],
    widths=[2.2, 4.6])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. FEATURE IMPORTANCE
# ══════════════════════════════════════════════════════════════
H1("4.  Feature Importance — Ranked & Interpreted")
para("XGBoost reports how often and how decisively each feature was used across all 500 trees. "
     "Below is the final ranking on the 16-feature model.")
imp = pd.read_csv("outputs/feature_importance.csv")
rows = [[i+1, r["Feature"], f"{r['Importance_pct']:.2f}%"] for i, r in imp.iterrows()]
make_table(["Rank", "Feature", "Importance %"], rows, widths=[0.8, 3.5, 1.5])

H3("How to Read This")
bullet("lag_1 (62%) + lag_4 (30%) dominate — recent store performance is the backbone of any short-horizon sales forecast. This is expected and correct for week-ahead prediction.", )
bullet("After lags, the leading signals are Wallflowers Depot (competition), housing_new_share, CYE Total Population, and CYE Median Household Income — the market features earn a meaningful, interpretable share.")
bullet("Why this matters: planners distrust black boxes. By trimming the rolling averages we made the model's reasoning legible — 'recent sales, adjusted for competition, housing growth, income, and population' — which is exactly the story a planner can defend in a target review.")
para("Important nuance: feature importance is for 1-week-ahead prediction, where recent sales "
     "already encode slow-moving market effects. For a brand-new store (no lag history) or a "
     "multi-year forecast, the demographic and competition features would carry far more weight — "
     "which is why we kept them in the model.", italic=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 5. METHODOLOGY & INTEGRITY
# ══════════════════════════════════════════════════════════════
H1("5.  Methodology & Model Integrity")

H2("5.1  Leakage Prevention (the most important guard)")
para("Three columns are permanently blacklisted as inputs because they are same-period measures "
     "that would not be available when the prediction is actually needed:")
make_table(["Forbidden Column", "Corr w/ Target", "Why It's Leakage"],
    [["Plan Sales USD", "0.986", "It IS essentially the answer — the plan we are trying to improve upon."],
     ["Invoice Count", "0.972", "Same-period transaction count; Sales = Invoices × Avg Ticket."],
     ["Avg Ticket", "0.133", "Same-period decomposition of Sales."]],
    widths=[1.8, 1.4, 3.6])
para("A programmatic assertion in feature_engineering.py raises an error if any blacklisted column "
     "ever enters the feature list. All lag features use shift(1) minimum.")

H2("5.2  Time-Based Validation (not random)")
bullet("Rolling split, NOT random — random splitting on time series leaks the future into training.")
bullet("Train: all data before 19 Sep 2025. Holdout: 19 Sep 2025 → 30 Jan 2026 (Q4 FY2025).")
bullet("22,439 holdout rows spanning all 1,727 stores → the model only predicts the future from the past.")

H2("5.3  Multicollinearity Handling")
para("Raw income brackets (8 columns) and housing decades (9 columns) each sum to ~100%. Feeding "
     "all of them in creates redundant, unstable inputs. We engineered strategic combinations "
     "(income_affluent, housing_new_share, housing_old_share) that preserve the signal without "
     "the instability.")

H2("5.4  The Store-Level Accuracy Loop")
para("Instead of trusting one global WAPE (which can hide offsetting errors), we compute per-store "
     "WAPE, Bias, RMSE and quarterly bias, then classify each store:")
make_table(["Classification", "Rule", "Planner Action"],
    [["Under-Targeted", "bias < -5%", "Store sandbagged — raise the target."],
     ["Over-Targeted", "bias > +5%", "Target too aggressive — review for headwinds."],
     ["Well-Calibrated", "within ±5%", "Trust the model output directly."]],
    widths=[1.7, 1.3, 3.8])
para("Result on the final model: 1,636 well-calibrated, 60 under-targeted, 31 over-targeted. "
     "This enables surgical overrides instead of chain-wide guesswork — the core business value.")

H2("5.5  Role-of-Store Segmentation")
para("A post-prediction overlay. We compute a growth_score (CAGR households, new-housing share, "
     "minus model bias) and a risk_score (competitive pressure plus bias), take the net, and split "
     "stores into quintiles → High Growth, Growth, Neutral, Maintain, Defend (≈20% each). "
     "Quantile bucketing was chosen after fixed thresholds produced an unusable 86% 'High Growth'. "
     "Each role becomes a distinct planning playbook.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 6. RESULTS & BUSINESS IMPACT
# ══════════════════════════════════════════════════════════════
H1("6.  Results & Business Impact")
bullet("Accuracy: 7.05% WAPE vs 8.27% naive — a 15% reduction in error, on 1,727 stores.", bold_prefix="Accuracy")
bullet("Bias: -0.97% means the model neither systematically over- nor under-commits — the annual plan reconciles cleanly.", bold_prefix="Bias")
bullet("Calibration: 94.7% of stores within ±5% bias — planners can trust the vast majority of targets out-of-the-box.", bold_prefix="Calibration")
bullet("Targeted overrides: only ~91 stores need human review (vs. chain-wide manual replanning today).", bold_prefix="Efficiency")
bullet("Explainability: a 16-feature model whose drivers a planner can read and defend.", bold_prefix="Trust")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 7. WHAT WORKED WELL / REFLECTIONS
# ══════════════════════════════════════════════════════════════
H1("7.  What Worked Well — Reflections")
H3("What we liked")
bullet("The lag-feature design. Engineering store-grouped, leakage-safe lags turned a hard forecasting problem into a tractable one — they alone explain ~90% of predictive power.")
bullet("Trimming to 16 features. Counter-intuitively, removing the two strongest rolling features improved bias and made the model presentable. Fewer, well-chosen features beat a kitchen-sink model.")
bullet("The Store-Level Accuracy Loop. This is the real differentiator versus a generic forecast — it converts model output into a concrete, prioritized planner action list.")
bullet("Model agreement. XGBoost, LightGBM and Ridge all landing near 7% WAPE gave high confidence the result is real and robust, not an artifact.")
bullet("Discipline on leakage. Refusing to use Plan Sales USD (corr 0.986) cost headline accuracy but produced a model that actually works in production — the right trade-off.")

H3("Honest limitations")
bullet("Lag dominance means the model is strongest for existing stores with history; brand-new stores need a cold-start variant relying on demographics only.")
bullet("Pure week-ahead scoring. True forward forecasting of an unseen future year needs recursive prediction (predict week 1 → feed as lag → predict week 2). A clear next step.")
bullet("Role-of-Store thresholds are heuristic. A future iteration could derive them from a proper clustering model.")
bullet("Demographics are annual snapshots; intra-year shifts are not captured.")

H3("Recommended next steps")
bullet("Add recursive multi-week forecasting for true next-FY target setting.")
bullet("Build a cold-start (no-lag) model for new stores using only market features.")
bullet("Layer constrained optimization so store targets sum exactly to the division plan.")
bullet("Pilot on the West Division, measure override-rate reduction vs. the current process, then scale chain-wide.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 8. DELIVERABLES & HOW TO RUN
# ══════════════════════════════════════════════════════════════
H1("8.  Deliverables & How to Run")
make_table(["Artifact", "Location"],
    [["ML pipeline (feature eng., training, store loop)", "src/*.py"],
     ["Interactive presentation app", "app/streamlit_app.py"],
     ["Trained model + results", "outputs/ (pkl + CSVs)"],
     ["Feature documentation (Excel)", "outputs/Phase1_Feature_Documentation.xlsx"],
     ["This document", "outputs/Capstone8_Full_Documentation.docx"],
     ["Source repository", "github.com/nishthaspeakx/capstone_8"]],
    widths=[3.4, 3.4])
H3("Reproduce end-to-end")
para("pip install -r requirements.txt", )
para("python src/train_model.py        # retrains, regenerates outputs/")
para("streamlit run app/streamlit_app.py   # launches the dashboard")

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Capstone 8 · Group 8 · IIM Calcutta APAL02 · Lowe's Store-Level Sales Target Model")
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY

out = "outputs/Capstone8_Full_Documentation.docx"
doc.save(out)
desktop = os.path.expanduser("~/Desktop/Capstone8_Full_Documentation.docx")
shutil.copy(out, desktop)
print(f"Saved: {out}")
print(f"Desktop copy: {desktop}")
print(f"Size: {os.path.getsize(out)/1024:.1f} KB")
