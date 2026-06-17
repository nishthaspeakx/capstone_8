"""
Build the LIVE DEMO speaker notes PDF for the Streamlit walkthrough.
Output: outputs/Capstone8_LiveDemo_SpeakerNotes.pdf  (also copied to Desktop)
"""
import os, shutil, subprocess
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x00, 0x30, 0x87)
BLUE = RGBColor(0x02, 0x77, 0xBD)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
ORANGE = RGBColor(0xE6, 0x51, 0x00)
GREY = RGBColor(0x55, 0x55, 0x55)
ACCENT = RGBColor(0xF4, 0xB4, 0x00)
INK = RGBColor(0x1A, 0x23, 0x32)

doc = Document()
for section in doc.sections:
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def H1(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.color.rgb = NAVY; r.font.size = Pt(20); r.bold = True
    return p

def H2(text, color=BLUE):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.color.rgb = color; r.font.size = Pt(13); r.bold = True
    return p

def label(p, text, color, size=10, bold=True):
    r = p.add_run(text); r.font.color.rgb = color; r.bold = bold; r.font.size = Pt(size)
    return r

def line(text, size=11, color=None, italic=False, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(size); r.italic = italic; r.bold = bold
    if color: r.font.color.rgb = color
    return p

def spacer(pt=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pt)
    return p

# ════════════════════════════════════════════════════════════
# COVER
# ════════════════════════════════════════════════════════════
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("LIVE DEMO — Speaker Notes")
r.font.size = Pt(22); r.bold = True; r.font.color.rgb = NAVY

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Lowe's Store-Level AI Sales Targeting · Streamlit Walkthrough")
r.font.size = Pt(12); r.italic = True; r.font.color.rgb = BLUE

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Capstone 8 · Group 8 · IIM Calcutta APAL02   |   Target duration: ~4 minutes")
r.font.size = Pt(10); r.font.color.rgb = GREY

spacer(8)

# Mini overview
t = doc.add_table(rows=1, cols=2); t.style = "Light Grid Accent 1"
hdr = t.rows[0].cells
for c, txt in zip(hdr, ["Section", "Time"]):
    c.text = ""; rr = c.paragraphs[0].add_run(txt); rr.bold = True; rr.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    shade(c, "003087")
overview = [
    ("0 · Opening + the hero stat banner", "0:00 – 0:20"),
    ("1 · Pick-a-Store simulator — Best Save (Wolfridge, TX)", "0:20 – 1:15"),
    ("2 · Click EXPLAIN — the GenAI explainer", "1:15 – 1:50"),
    ("3 · Switch to an UNDER-TARGETED store", "1:50 – 2:20"),
    ("4 · Scroll to RESULTS + comparison table", "2:20 – 2:45"),
    ("5 · Feature importance chart", "2:45 – 3:05"),
    ("6 · Store-Level Accuracy Loop (the differentiator)", "3:05 – 3:35"),
    ("7 · Role-of-Store doughnut", "3:35 – 3:55"),
    ("8 · Close + hand back to the deck", "3:55 – 4:10"),
]
for sec, time_ in overview:
    row = t.add_row().cells
    row[0].text = ""; row[0].paragraphs[0].add_run(sec).font.size = Pt(10)
    row[1].text = ""; row[1].paragraphs[0].add_run(time_).font.size = Pt(10)
    row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
for row in t.rows:
    row.cells[0].width = Inches(5.7); row.cells[1].width = Inches(1.5)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# Generic helper: section
# ════════════════════════════════════════════════════════════
def section(title_text, window, do_lines, say_text, transition_text=None, callout=None):
    H1(title_text)
    line(window, size=11, color=GREY, italic=True)
    spacer(2)

    # DO block
    p = doc.add_paragraph(); label(p, "DO  ", ACCENT, size=11, bold=True)
    p.paragraph_format.space_after = Pt(2)
    for d in do_lines:
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_after = Pt(2)
        rr = bp.add_run(d); rr.font.size = Pt(11)
    spacer(2)

    # SAY block
    p = doc.add_paragraph(); label(p, "SAY  ", GREEN, size=11, bold=True)
    p.paragraph_format.space_after = Pt(2)
    sp = doc.add_paragraph()
    sp.paragraph_format.left_indent = Inches(0.25)
    sp.paragraph_format.space_after = Pt(4)
    rr = sp.add_run('"' + say_text + '"'); rr.font.size = Pt(12); rr.italic = True

    if callout:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        rr = p.add_run("Tip: " + callout); rr.font.size = Pt(9.5); rr.font.color.rgb = ORANGE; rr.italic = True

    if transition_text:
        p = doc.add_paragraph()
        label(p, "NEXT  ", BLUE, size=10, bold=True)
        rr = p.add_run(transition_text); rr.font.size = Pt(10.5); rr.italic = True

    spacer(8)

# ════════════════════════════════════════════════════════════
# SECTION 0 — Opening
# ════════════════════════════════════════════════════════════
section(
    "0 · Opening + Hero Stat Banner",
    "0:00 – 0:20  (~20 seconds)",
    [
        "Switch to the browser tab with the Streamlit app open (http://localhost:8501 or the deployed URL).",
        "Make sure the page is scrolled to the very top so the hero banner is visible.",
    ],
    "Now let me show you the dashboard that the planner actually uses. The first thing you see is the headline outcome — 94.7% of stores get accurate sales targets within plus-or-minus five percent. The override workload drops from 1,727 stores every cycle today, to only 91 stores with our model — a 95% reduction.",
    transition_text="Scroll down to the 'Try It: Pick a Store' section.",
    callout="Have the app pre-loaded BEFORE the demo starts. Don't switch tabs live — keep it ready."
)

# ════════════════════════════════════════════════════════════
# SECTION 1 — Pick-a-Store Best Save
# ════════════════════════════════════════════════════════════
section(
    "1 · Pick-a-Store Simulator — Best Save",
    "0:20 – 1:15  (~55 seconds)",
    [
        "In the dropdown, type 'S01155' or 'Wolfridge' to find the best-save example.",
        "Wait for the three side-by-side comparison cards to render.",
    ],
    "This is where the model meets reality. We pick a real store — Wolfridge, Texas. Here are three numbers: the actual sales for the holdout period — about 5 million dollars. The current peanut-butter target — 6.5 million, off by 30%. And our model's target — 5.14 million, off by just 7%. The model fixes a 23-percentage-point error. And notice the green verdict callout below: 'Model wins — 23 pp improvement for this store.' That is what data-driven planning looks like on a single store.",
    transition_text="Stay on the same store. Now click the 'Explain' button.",
    callout="If S01155 is busy, any 'Best Save' store works — sort by Model_Improvement_pp in the loop CSV."
)

# ════════════════════════════════════════════════════════════
# SECTION 2 — Click EXPLAIN
# ════════════════════════════════════════════════════════════
section(
    "2 · Click EXPLAIN — the GenAI Explainer",
    "1:15 – 1:50  (~35 seconds)",
    [
        "Click the blue 'Explain' button next to the 'Why did this store get this target?' heading.",
        "Wait ~2-3 seconds for the explanation card to appear.",
        "Read the first two sentences aloud to ground the professor in the output.",
    ],
    "This is the GenAI layer the model produces a plain-English explanation for any store. It references the actual numbers — the recent run-rate, the year-over-year change, competitor count, housing share — and tells the planner exactly why the target landed where it did and what action to take: trust it, raise it, or review it. Every one of our 1,727 stores has this on demand.",
    transition_text="Open the dropdown again and pick an under-targeted store.",
    callout="If Gemini quota hits a limit, the deterministic fallback still produces a clean explanation. Either way looks identical to the audience."
)

# ════════════════════════════════════════════════════════════
# SECTION 3 — Switch to Under-Targeted
# ════════════════════════════════════════════════════════════
section(
    "3 · Switch to an Under-Targeted Store",
    "1:50 – 2:20  (~30 seconds)",
    [
        "From the dropdown, pick 'S01442 — Blackreach, GA' (largest under-targeting case).",
        "Point at the three comparison cards, then at the orange 'Plan wins on this one' verdict.",
        "Optionally click Explain to refresh the per-store narrative.",
    ],
    "Here is the opposite case — Blackreach, Georgia. Actual sales blew past both the peanut-butter target and our model's prediction. Our model still under-predicts by about 22%, but it tells the planner exactly which stores it can't fully explain — that's the orange warning. These ~91 stores are the only ones a planner needs to review. The other 1,636 — the green ones — they can just trust.",
    transition_text="Now scroll down past the simulator to 'Headline Results'.",
    callout="The point: model is honest about what it can't explain — that builds planner trust."
)

# ════════════════════════════════════════════════════════════
# SECTION 4 — Results
# ════════════════════════════════════════════════════════════
section(
    "4 · Headline Results + Comparison Table",
    "2:20 – 2:45  (~25 seconds)",
    [
        "Scroll to '#2 · Headline Results'. Five KPI cards should be in view.",
        "Point at WAPE 7.05% (green check), Bias −0.97% (green check), 94.7% calibrated (gold).",
        "Briefly point to the comparison table below.",
    ],
    "Every success criterion is met — WAPE under 8%, bias within ±2%, R-squared above 0.88. And as you can see in this comparison table, all three real models — XGBoost, LightGBM, Ridge — agree at about 7%, while the naive baseline is at 8.27. That agreement tells us the result is robust, not a fluke of one algorithm.",
    transition_text="Scroll to the feature importance chart.",
)

# ════════════════════════════════════════════════════════════
# SECTION 5 — Feature Importance
# ════════════════════════════════════════════════════════════
section(
    "5 · Feature Importance",
    "2:45 – 3:05  (~20 seconds)",
    [
        "Scroll to '#4 · What Drives the Predictions' (the bar chart).",
        "Point at lag_1 and lag_4 at the top, then move down to housing_new_share / Total Population / Median Income.",
    ],
    "Recent sales — lag_1 and lag_4 — dominate, which is expected for a one-week-ahead forecast. But notice that after the lags, the next signals are competition (Wallflowers Depot), new housing share, total population, and median household income. Those are exactly the local market factors the peanut-butter spread completely ignores. That's why a planner can defend our targets — not a black box.",
    transition_text="Scroll to the Store-Level Accuracy Loop section.",
)

# ════════════════════════════════════════════════════════════
# SECTION 6 — Store Accuracy Loop
# ════════════════════════════════════════════════════════════
section(
    "6 · Store-Level Accuracy Loop",
    "3:05 – 3:35  (~30 seconds)",
    [
        "Scroll to '#5 · Store-Level Accuracy Loop'.",
        "Point at the 4 summary metrics: total, well-calibrated, under-, over-.",
        "Click the 'Under-Targeted' tab. Wait for the table to render. Show top 5 rows.",
        "Optional: click 'Over-Targeted' tab briefly.",
    ],
    "And this is our real differentiator. Per-store WAPE, Bias, quarterly bias — everything diagnosed. Sixty stores are under-targeted, thirty-one over-targeted. Those 91 are the planner's prioritized list. Everything else — 1,636 stores — is well-calibrated and ready to use as-is. This converts model output into a concrete, prioritized action queue for the planning team.",
    transition_text="Scroll to the Role-of-Store doughnut chart.",
)

# ════════════════════════════════════════════════════════════
# SECTION 7 — Role of Store
# ════════════════════════════════════════════════════════════
section(
    "7 · Role-of-Store Doughnut",
    "3:35 – 3:55  (~20 seconds)",
    [
        "Scroll to '#6 · Role-of-Store Segmentation'. Doughnut chart visible.",
        "Briefly point at each segment color in turn.",
    ],
    "And the strategic layer on top — Role-of-Store. Every store falls into one of five segments — High Growth, Growth, Neutral, Maintain, Defend — derived from household growth, new-housing share, and competitive intensity. So one model becomes five distinct planning playbooks, not just a number-per-store.",
    transition_text="Scroll back to the top of the page to close cleanly.",
)

# ════════════════════════════════════════════════════════════
# SECTION 8 — Close
# ════════════════════════════════════════════════════════════
section(
    "8 · Close + Hand Back to the Deck",
    "3:55 – 4:10  (~15 seconds)",
    [
        "Scroll back to the top of the dashboard so the 94.7% hero banner is visible.",
        "Switch back to the PowerPoint deck (Cmd+Tab → Keynote/PowerPoint), advance to the Thank-You slide.",
    ],
    "So that's the full system — model on the back end, plain-English explanations for every store, a prioritized override queue, and strategic segmentation. We replaced one flat number with 1,727 store-specific targets that planners can actually trust and defend. Happy to take questions.",
    transition_text="The deck's last slide is the recap + Q&A landing card.",
)

# ════════════════════════════════════════════════════════════
# Q&A PREP (last page)
# ════════════════════════════════════════════════════════════
doc.add_page_break()
H1("Likely Q&A — quick answers")
spacer(4)

qas = [
    ("Why XGBoost over LightGBM?", "Best WAPE (7.05 vs 7.07) AND best bias (-0.97 vs -1.11). Both are gradient boosted — LightGBM agreeing confirms robustness."),
    ("How do you know it's not overfitting?", "Train WAPE 6.62 vs test WAPE 7.05 — a 0.4 pp gap. Time-based holdout (Sep 19, 2025 onward), never random."),
    ("Why did you DROP feature columns to 16 from 40?", "Cutting roll_4, roll_13, Year & Fiscal Week reduced lag dominance from 94% to 91%, made bias near-zero, and let market features show their real signal — model became more explainable AND better-calibrated."),
    ("What about new stores with no lag history?", "Honest limitation — current model is lag-driven. Next step is a cold-start variant using only demographics, housing, and competition."),
    ("Why not use Plan Sales USD as an input?", "Correlation 0.986 — using it gives fake 99% R² but no real-world value, because at prediction time you don't HAVE next week's plan — that IS what the model is supposed to generate."),
    ("How does this scale to all 1,727 stores chain-wide?", "Already does — every result you saw covers all 1,727 stores. Phase 2 pilots on West Division to measure override-rate reduction before chain-wide rollout."),
    ("Is the GenAI explanation reliable?", "Two layers: Gemini-generated primary, deterministic rule-based fallback. Both reference the same 16 feature values — same accuracy, just different prose."),
    ("How does this compare to the current peanut-butter spread?", "Our model: 94.7% of stores within ±5%. Current method: 74% of store-weeks miss by >10%. About 95% reduction in planner override workload."),
]
for q, a in qas:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    rr = p.add_run("Q: "); rr.bold = True; rr.font.color.rgb = NAVY; rr.font.size = Pt(11)
    rr = p.add_run(q); rr.bold = True; rr.font.size = Pt(11)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    rr = p.add_run("A: "); rr.bold = True; rr.font.color.rgb = GREEN; rr.font.size = Pt(11)
    rr = p.add_run(a); rr.font.size = Pt(10.5)

# Footer note
spacer(6)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Capstone 8 · Group 8 · IIM Calcutta APAL02   |   storewisetarget.streamlit.app   |   github.com/nishthaspeakx/capstone_8")
r.font.size = Pt(9); r.font.color.rgb = GREY; r.italic = True

# Save
docx_path = "outputs/Capstone8_LiveDemo_SpeakerNotes.docx"
doc.save(docx_path)
print(f"Saved DOCX: {docx_path}")

# Convert to PDF via LibreOffice
print("Converting to PDF…")
result = subprocess.run(
    ["soffice","--headless","--convert-to","pdf",docx_path,"--outdir","outputs"],
    capture_output=True, text=True, timeout=120)
print(result.stdout[-200:] if result.stdout else result.stderr[-300:])

pdf_path = docx_path.replace(".docx",".pdf")
desktop_pdf = os.path.expanduser("~/Desktop/Capstone8_LiveDemo_SpeakerNotes.pdf")
shutil.copy(pdf_path, desktop_pdf)
print(f"Saved PDF: {pdf_path}")
print(f"Desktop:   {desktop_pdf}")
print(f"Size: {os.path.getsize(pdf_path)/1024:.1f} KB")
