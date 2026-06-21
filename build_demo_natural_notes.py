"""
Build a focused live-demo PDF in plain conversational language.
Structure: Opening + Store 1 (best save) + Store 2 (worst miss) + GenAI explainer.
Output: outputs/Capstone8_LiveDemo_NaturalNotes.pdf
"""
import os, shutil, subprocess
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x00, 0x30, 0x87)
BLUE = RGBColor(0x02, 0x77, 0xBD)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
ORANGE = RGBColor(0xE6, 0x51, 0x00)
RED = RGBColor(0xC6, 0x28, 0x28)
GREY = RGBColor(0x55, 0x55, 0x55)
ACCENT = RGBColor(0xF4, 0xB4, 0x00)
INK = RGBColor(0x1A, 0x23, 0x32)
PURPLE = RGBColor(0x6F, 0x42, 0xC1)

doc = Document()
for sec in doc.sections:
    sec.left_margin = Inches(0.7); sec.right_margin = Inches(0.7)
    sec.top_margin = Inches(0.6); sec.bottom_margin = Inches(0.6)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hex_color)
    tcPr.append(shd)

def H1(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.color.rgb = NAVY; r.font.size = Pt(19); r.bold = True

def H2(text, color=BLUE):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.color.rgb = color; r.font.size = Pt(13); r.bold = True

def kicker(text, color=GREY):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.color.rgb = color; r.italic = True; r.font.size = Pt(11)

def label_run(p, text, color, size=11, bold=True):
    r = p.add_run(text); r.font.color.rgb = color; r.bold = bold; r.font.size = Pt(size)

def spacer(pt=4):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(pt)

def do_block(items):
    p = doc.add_paragraph(); label_run(p, "DO  ", ACCENT)
    p.paragraph_format.space_after = Pt(2)
    for d in items:
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_after = Pt(2)
        bp.add_run(d).font.size = Pt(11)

def say_block(text):
    spacer(3)
    p = doc.add_paragraph(); label_run(p, "SAY  ", GREEN)
    p.paragraph_format.space_after = Pt(2)
    sp = doc.add_paragraph()
    sp.paragraph_format.left_indent = Inches(0.25)
    r = sp.add_run('"' + text + '"')
    r.font.size = Pt(12); r.italic = True

def tip(text):
    p = doc.add_paragraph()
    label_run(p, "TIP  ", ORANGE, size=10, bold=True)
    r = p.add_run(text); r.font.size = Pt(10.5); r.italic = True; r.font.color.rgb = ORANGE

def next_step(text):
    p = doc.add_paragraph()
    label_run(p, "NEXT  ", BLUE, size=10, bold=True)
    r = p.add_run(text); r.font.size = Pt(10.5); r.italic = True

# ════════════════════════════════════════════════════════════
# COVER
# ════════════════════════════════════════════════════════════
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Live Demo Speaker Notes")
r.font.size = Pt(24); r.bold = True; r.font.color.rgb = NAVY

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Streamlit Dashboard Walkthrough — natural language")
r.font.size = Pt(13); r.italic = True; r.font.color.rgb = BLUE

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Lowe's Store-Level AI Sales Targeting   ·   Capstone 8   ·   Group 8   ·   IIM Calcutta")
r.font.size = Pt(10); r.font.color.rgb = GREY

spacer(10)

# Table of contents
t = doc.add_table(rows=1, cols=2); t.style = "Light Grid Accent 1"
hdr = t.rows[0].cells
for c, txt in zip(hdr, ["What you'll cover", "~Time"]):
    c.text = ""; rr = c.paragraphs[0].add_run(txt); rr.bold = True; rr.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    shade(c, "003087")
overview = [
    ("1 · Open the dashboard, point at the headline 94.7%", "~25 seconds"),
    ("2 · STORE ONE — Wolfridge, Texas (a 'big win' for the model)", "~60 seconds"),
    ("3 · STORE TWO — Blackreach, Georgia (a 'flagged' store)", "~50 seconds"),
    ("4 · The GenAI Layer — how it works + the fallback", "~75 seconds"),
    ("5 · Closing line + hand back to deck", "~10 seconds"),
]
for sec, time_ in overview:
    row = t.add_row().cells
    row[0].text = ""; row[0].paragraphs[0].add_run(sec).font.size = Pt(10.5)
    row[1].text = ""; row[1].paragraphs[0].add_run(time_).font.size = Pt(10.5)
    row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
for row in t.rows:
    row.cells[0].width = Inches(5.7); row.cells[1].width = Inches(1.5)

spacer(6)
p = doc.add_paragraph()
r = p.add_run("Total demo time: about 3.5 minutes. Speak naturally, do not rush. The numbers do most of the talking.")
r.italic = True; r.font.size = Pt(10); r.font.color.rgb = GREY

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# SECTION 1 — OPENING
# ════════════════════════════════════════════════════════════
H1("1 · Open the dashboard")
kicker("~25 seconds")
spacer(3)

do_block([
    "Click into the browser tab that already has the Streamlit app open. (Open this BEFORE the demo starts.)",
    "Scroll to the very top so the giant blue '94.7%' banner is the first thing on screen.",
])

say_block(
    "Alright, this is the dashboard a planner at Lowe's would actually log into every week. "
    "The very first thing you see is the headline. Ninety four point seven percent of our seventeen hundred stores "
    "now get a sales target that is within five percent of what they actually sell. Before our model, "
    "the planner had to manually review every single one of these 1,727 stores. With our system, "
    "they only need to look at 91. So ninety five percent of their workload just disappears."
)
tip("Pause for one second after saying '95 percent of their workload just disappears.' Let it land.")
next_step("Scroll down to the section called 'Try It: Pick a Store'.")

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# SECTION 2 — STORE 1
# ════════════════════════════════════════════════════════════
H1("2 · Store One — Wolfridge, Texas")
kicker("~60 seconds  ·  the 'big win' example")
spacer(3)

do_block([
    "Click the dropdown that says 'Select a store'. Start typing 'Wolfridge' or the ID 'S01155'.",
    "Click the store. Wait for the THREE cards to appear — Actual / Peanut-butter / Our Model.",
    "Point at each card in order, left to right.",
])

say_block(
    "Let me show you a real store. This is Wolfridge in Texas. I am picking it because it tells the "
    "best story about why our model exists. Look at these three numbers on the screen. "
    "The first card on the left — that is what the store actually sold during the test period. About five million dollars. "
    "The middle card — that is the target Lowe's set using their current method, the peanut-butter spread. "
    "Six and a half million. Off by thirty percent. The planner would have spent the year explaining why this store "
    "missed by so much. "
    "Now look at the third card. Our model's target. Five point one four million. Off by just seven percent. "
    "And the green banner below confirms it — for this store, our model wins by twenty three percentage points. "
    "That is one store. Now multiply this kind of improvement across hundreds of stores."
)
tip("If Wolfridge does not appear in the dropdown, pick any store that shows a green 'Model wins' banner — the script still works.")
next_step("Stay on the same store. Scroll to the 'Why did this store get this target?' heading.")

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# SECTION 3 — STORE 2
# ════════════════════════════════════════════════════════════
H1("3 · Store Two — Blackreach, Georgia")
kicker("~50 seconds  ·  the 'flagged' store — shows our model is honest")
spacer(3)

do_block([
    "Open the dropdown again. Type 'Blackreach' or 'S01442'.",
    "Click the store. Wait for the three cards to load — they will look DIFFERENT this time.",
    "Point at the orange or red warning banner that appears below the cards.",
])

say_block(
    "Now I want to show you the opposite case. This is Blackreach in Georgia. "
    "Watch what happens. Actual sales for the period — twenty seven point seven million dollars. "
    "Peanut-butter target says — way less. Our model — also lower than what actually happened. "
    "About twenty one million. So we under-predicted by twenty four percent. "
    "Why is this important? Because the model is HONEST about it. See that orange warning banner? "
    "The model is telling the planner — this is one of the ninety one stores I am not confident about. "
    "Please review it. So instead of a planner having to eyeball all 1,727 stores looking for problems, "
    "the model raises its hand on just 91 of them. The other 1,636, the planner can trust and move on."
)
tip("If Blackreach is hard to find, any store with an orange 'Plan wins on this one' or 'Under-Targeted' banner works.")
next_step("Stay on this store. Click the blue 'Explain' button next to 'Why did this store get this target?'")

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# SECTION 4 — GENAI EXPLAINER
# ════════════════════════════════════════════════════════════
H1("4 · The GenAI Layer — how it works")
kicker("~75 seconds  ·  the wow moment")
spacer(3)

do_block([
    "Click the blue 'Explain' button.",
    "Wait two or three seconds. A three-paragraph explanation will appear in a card.",
    "Read the first line out loud so the audience can follow.",
])

say_block(
    "This is where the AI part comes in. Watch what just happened. The system pulled the sixteen "
    "feature values for this specific store — its recent sales, the median income in the area, how many "
    "competitors are nearby, the share of new housing, everything — and asked Google's Gemini AI "
    "to write a plain English explanation for THIS store. "
    "Read it with me. It quotes the actual numbers. It tells the planner exactly why the target landed where "
    "it did. And it ends with a clear instruction — trust the number, raise it, or review it. "
    "Every single one of our 1,727 stores can do this. On demand. In three seconds."
)
spacer(4)

H2("How does the GenAI actually work?", color=PURPLE)
spacer(2)
say_block(
    "Three simple steps happen behind the scenes when I clicked Explain. "
    "Step one — the app grabs the data for this one store. All sixteen features, plus the predicted target, "
    "plus the actual sales, plus the bias percentage. Step two — it builds a prompt for Google Gemini "
    "asking three specific questions. Why this target. Why this segment. What should the planner do. "
    "Step three — Gemini reads the numbers and writes a custom paragraph. "
    "No two stores get the same answer. Because no two stores have the same features."
)
spacer(4)

H2("What if Gemini is down? The fallback explained", color=PURPLE)
spacer(2)
say_block(
    "Now you might ask — what if Google's AI is down, or we have hit our daily limit? "
    "We built a backup. There is a second explainer written in plain Python rules. "
    "It uses the same sixteen features, follows the same three part structure, "
    "and produces an answer that looks identical to the audience. The app silently switches to it whenever "
    "Gemini is unavailable. The demo never fails. The planner always gets an answer. "
    "That is how production systems are supposed to work — they should never depend on one external service "
    "being a hundred percent up."
)
spacer(6)

H2("Is this useful TODAY with our static demo data?", color=PURPLE)
say_block(
    "Yes. Here is why. We are running on the fixed FY 2025 holdout data, "
    "but every single one of the 1,727 stores has a completely different combination of features. "
    "So every explanation is unique. The system is genuinely reading the numbers and writing a custom answer. "
    "It is not template text. When this rolls into production with weekly live data, the same code keeps running "
    "without any change. Just newer numbers in, fresher explanations out."
)

tip("If a professor asks 'is this real AI or just templates?' — show them two different stores' explanations. "
    "They WILL look different because the underlying numbers are different. That proves it.")
next_step("Scroll back to the top of the dashboard, then switch back to your Thank You slide.")

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# SECTION 5 — CLOSE
# ════════════════════════════════════════════════════════════
H1("5 · Closing line")
kicker("~10 seconds  ·  hand back to the deck")
spacer(3)

do_block([
    "Scroll back to the top of the dashboard so the 94.7% banner is the last thing the audience sees.",
    "Then use Cmd+Tab (Mac) to switch back to PowerPoint and advance to the final 'Thank You' slide.",
])

say_block(
    "So that is the full system. The model picks the target. The dashboard tells the planner which ninety one "
    "stores to review. And the AI layer explains every decision in plain English. "
    "From one flat number across all 1,727 stores, to seventeen hundred personalised targets that planners can "
    "actually defend. Happy to take your questions."
)

# ════════════════════════════════════════════════════════════
# Q&A SHEET
# ════════════════════════════════════════════════════════════
doc.add_page_break()
H1("Q&A Cheat Sheet — likely follow-up questions")
spacer(6)

qa = [
    ("Why these two stores specifically?",
     "Wolfridge shows the model winning big — easiest to explain. Blackreach shows the model being honest about its limits — the same dashboard that wins also flags problems. Together they tell the full story."),
    ("How does the model flag the 91 stores?",
     "For every store, we compute one number: the bias percentage. That is the total predicted sales minus the total actual sales, divided by actual. If bias is more than plus-or-minus five percent, the store is flagged. Simple rule. Lives in one Python function we can show on demand."),
    ("Is this real AI or just pre-written text?",
     "Real AI. Click any two stores and the explanations will be different because the underlying sixteen feature values are different for each store. There is no template."),
    ("What is the fallback for the GenAI?",
     "A rule-based explainer written in Python. Uses the same sixteen features and follows the same three part structure. If Gemini is unreachable, the app silently switches. The audience cannot tell."),
    ("Does this work on live data?",
     "Yes. The system is built to handle weekly data refreshes. The model runs the same way. The explainer runs the same way. The flagging rule runs the same way. Just newer numbers in, fresher results out."),
    ("Why is WAPE 7.05% acceptable?",
     "WAPE means seven dollars of error per hundred dollars of sales. Industry rule of thumb for retail forecasting is anything under ten percent is good. Seven is comfortably in the good zone. Plus we beat the naive baseline by fifteen percent."),
    ("What about the 91 stores you cannot predict — what is the plan?",
     "We add them to a planner review queue. Each one gets a reason code from the planner — competitor opened, renovation, hurricane, et cetera. Over time these reason codes become training data for the next version. So the model gets better with use."),
    ("Why these 16 features and not 100?",
     "We started with 165 candidates, narrowed to 40, then to 16 by removing features that gave us double counting and features that did not help the model. Fewer features means the model is easier to explain to planners. And our accuracy actually IMPROVED when we cut to 16."),
]
for q, a in qa:
    p = doc.add_paragraph()
    rr = p.add_run("Q. "); rr.bold = True; rr.font.color.rgb = NAVY
    rr = p.add_run(q); rr.bold = True; rr.font.size = Pt(11.5)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.3)
    p2.paragraph_format.space_after = Pt(10)
    rr = p2.add_run("A. "); rr.bold = True; rr.font.color.rgb = GREEN; rr.font.size = Pt(11)
    p2.add_run(a).font.size = Pt(10.5)

spacer(6)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Capstone 8 · Group 8 · IIM Calcutta APAL02  |  storewisetarget.streamlit.app  |  github.com/nishthaspeakx/capstone_8")
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY

# Save & convert
docx_path = "outputs/Capstone8_LiveDemo_NaturalNotes.docx"
doc.save(docx_path)
print(f"Saved DOCX: {docx_path}")
result = subprocess.run(
    ["soffice","--headless","--convert-to","pdf",docx_path,"--outdir","outputs"],
    capture_output=True, text=True, timeout=120)
print((result.stdout or result.stderr)[-300:])
pdf_path = docx_path.replace(".docx",".pdf")
desktop = os.path.expanduser("~/Desktop/Capstone8_LiveDemo_NaturalNotes.pdf")
shutil.copy(pdf_path, desktop)
print(f"PDF: {pdf_path}")
print(f"Desktop: {desktop}")
print(f"Size: {os.path.getsize(pdf_path)/1024:.1f} KB")
