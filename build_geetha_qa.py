"""
Build a detailed Q&A PDF answering Geetha's 7 review-board questions.
Output: outputs/Capstone8_Geetha_QA.pdf
"""
import os, shutil, subprocess
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY  = RGBColor(0x00, 0x30, 0x87)
BLUE  = RGBColor(0x02, 0x77, 0xBD)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
ORANGE = RGBColor(0xE6, 0x51, 0x00)
GREY  = RGBColor(0x55, 0x55, 0x55)
ACCENT = RGBColor(0xF4, 0xB4, 0x00)
INK   = RGBColor(0x1A, 0x23, 0x32)
PURPLE = RGBColor(0x6F, 0x42, 0xC1)

doc = Document()
for sec in doc.sections:
    sec.left_margin = Inches(0.7); sec.right_margin = Inches(0.7)
    sec.top_margin = Inches(0.6); sec.bottom_margin = Inches(0.6)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hex_color)
    tcPr.append(shd)

def H1(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.color.rgb = NAVY; r.font.size = Pt(17); r.bold = True
    return p

def H2(text, color=BLUE):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.color.rgb = color; r.font.size = Pt(13); r.bold = True
    return p

def H3(text, color=PURPLE):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.color.rgb = color; r.font.size = Pt(11.5); r.bold = True
    return p

def para(text, size=11, italic=False, bold=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(size); r.italic = italic; r.bold = bold
    if color: r.font.color.rgb = color
    return p

def bullet(text):
    bp = doc.add_paragraph(style="List Bullet")
    bp.paragraph_format.space_after = Pt(3)
    bp.add_run(text).font.size = Pt(11)
    return bp

def numbered(text):
    bp = doc.add_paragraph(style="List Number")
    bp.paragraph_format.space_after = Pt(3)
    bp.add_run(text).font.size = Pt(11)
    return bp

def spacer(pt=4):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(pt)

def make_table(headers, rows, widths=None, header_color="003087"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True; run.font.size = Pt(10.5); run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shade(hdr[i], header_color)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t

# ════════════════════════════════════════════════════════════
# COVER
# ════════════════════════════════════════════════════════════
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Talking-Point Q&A — Detailed Answers"); r.font.size = Pt(22); r.bold = True; r.font.color.rgb = NAVY
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Lowe's Store-Level AI Sales Targeting · Capstone 8 · Group 8 · IIM Calcutta APAL02")
r.font.size = Pt(11); r.italic = True; r.font.color.rgb = BLUE
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Prepared for review · Seven questions answered in technical depth")
r.font.size = Pt(10); r.font.color.rgb = GREY
spacer(8)

# Table of contents
t = doc.add_table(rows=1, cols=2); t.style = "Light Grid Accent 1"
hdr = t.rows[0].cells
for c, txt in zip(hdr, ["Question", "Page"]):
    c.text = ""; rr = c.paragraphs[0].add_run(txt); rr.bold = True; rr.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    shade(c, "003087")
toc = [
    ("Q1 · Feature engineering techniques used to decide features", "p. 2"),
    ("Q2 · Standardization & normalization techniques (and why)", "p. 4"),
    ("Q3 · Hyperparameters of each of the 4 models", "p. 5"),
    ("Q4 · Why WAPE? + the final 16 feature importance weights", "p. 7"),
    ("Q5 · How the 5 store-role groups are chosen + thresholds", "p. 9"),
    ("Q6 · Features ignored due to multicollinearity (and why)", "p. 11"),
    ("Q7 · 'subsample 0.8' & 'colsample 0.8' in plain English", "p. 13"),
]
for q, page in toc:
    row = t.add_row().cells
    row[0].text = ""; row[0].paragraphs[0].add_run(q).font.size = Pt(10.5)
    row[1].text = ""; row[1].paragraphs[0].add_run(page).font.size = Pt(10.5)
    row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
for row in t.rows:
    row.cells[0].width = Inches(5.6); row.cells[1].width = Inches(1.0)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# Q1 — FEATURE ENGINEERING
# ════════════════════════════════════════════════════════════
H1("Q1 · Feature engineering techniques used to decide features")
para("We started with 194 raw columns and narrowed to 16 final features across three iterative rounds. "
     "Seven distinct techniques drove every keep/drop decision.", italic=True, color=GREY)
spacer(4)

H2("The 7 techniques applied")

H3("1. Correlation analysis against target")
para("For every numeric column we computed Pearson correlation with Actual Sales USD. "
     "Features with very low absolute correlation (≈ 0) were candidates for removal — but only if no other "
     "technique (domain logic, interaction effects) justified keeping them.")
para("Strongest correlations found in lag features (lag_1 = 0.982, roll_4 = 0.981); weakest in single-decade "
     "housing buckets and most distance-band competitor counts (typically |r| < 0.05).", color=GREY, italic=True)
spacer(3)

H3("2. Multicollinearity check (correlation matrix + sum-to-100 detection)")
para("We computed pairwise correlations between candidate features and flagged any pair with |r| > 0.85. "
     "We also detected feature groups that summed to ~100% (income-bracket columns, housing-decade columns). "
     "These were collapsed into 2-3 strategic engineered features to avoid coefficient instability.")
spacer(3)

H3("3. Variance / sparsity threshold")
para("Features with near-zero variance — for instance, competitor counts that were 0 for >95% of stores — "
     "were dropped. These cannot help any model split because they don't separate observations.")
para("Examples dropped: Thin Lizzy Hardware (100% zeros), Bobs Supplies (99% zeros).", color=GREY, italic=True)
spacer(3)

H3("4. Lag-feature engineering (time-series shift operations)")
para("Per-store grouped shift to create temporal features:")
para("• lag_1, lag_4, lag_13, lag_52 = sales shifted by 1, 4, 13, and 52 weeks per store", size=10.5)
para("• roll_4, roll_13 = rolling-mean of past 4 / 13 weeks (later dropped — see Q6)", size=10.5)
para("All lag operations used shift(1) minimum to prevent leakage (the model can never see the week it predicts).", italic=True, color=GREY)
spacer(3)

H3("5. Aggregation engineering (combining redundant raw columns)")
para("To preserve signal but avoid multicollinearity, we engineered composite features:")
bullet("income_affluent = % HH $150K-250K + % HH $250K+ (collapses 2 income brackets)")
bullet("housing_new_share = % built 2000-2009 + % built 2009-2019 (collapses 2 decade buckets)")
bullet("housing_old_share = % built pre-1949 + 1950s + 1960s (collapses 3 decade buckets)")
bullet("total_competitor_ta = sum of all 17 competitor counts in trade area (collapses 17 columns to 1)")
spacer(3)

H3("6. Categorical encoding")
para("Three categorical features were converted to integer codes using sklearn LabelEncoder:")
make_table(["Original", "Encoded as", "Levels"],
    [["Urbanicity", "Urbanicity_enc", "0 = Remote, …, 6 = Metropolis (7 levels)"],
     ["CBSA Type", "CBSA Type_enc", "0 = Non-CBSA, 1 = Micro, 2 = Metro (3 levels)"],
     ["CBSA Metro Size", "CBSA Metro Size_enc", "4 ordinal levels"]],
    widths=[1.8, 1.8, 3.4])
para("LabelEncoder (not OneHotEncoder) chosen because tree-based models split on numeric thresholds and "
     "handle ordinal codes natively without inflating the feature count.", italic=True, color=GREY)
spacer(3)

H3("7. Leakage detection (the critical guard)")
para("Three columns had correlations dangerously close to 1.0 — they are 'same-period' measures that would "
     "give the model the answer instead of features it has to learn from:")
make_table(["Banned column", "Correlation w/ target", "Why it's leakage"],
    [["Plan Sales USD", "0.986", "It IS the target — what we're trying to predict"],
     ["Invoice Count", "0.972", "Same-period transaction count (Sales = Invoices × Avg Ticket)"],
     ["Avg Ticket", "0.133", "Same-period decomposition of Sales"]],
    widths=[2.0, 1.4, 3.6])
para("A programmatic assertion in feature_engineering.py raises an error if any of these ever enter the feature list.", italic=True, color=GREY)
spacer(4)

H2("The funnel")
make_table(["Iteration", "Count", "What changed"],
    [["Raw columns", "194", "All columns in source CSV"],
     ["After identity/leakage drop", "165", "Removed Plan Sales, Invoice Count, Avg Ticket, store names, IDs"],
     ["Phase 1 set (1st round)", "40", "Business-relevance + correlation + variance threshold"],
     ["After 1st model iteration", "—", "XGBoost feature importance + multicollinearity check"],
     ["Final 16 features", "16", "Removed dominant lag-only features + redundant aggregates"]],
    widths=[2.0, 0.8, 4.2])

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# Q2 — STANDARDIZATION / NORMALIZATION
# ════════════════════════════════════════════════════════════
H1("Q2 · Standardization & normalization techniques (and why)")
spacer(4)

H2("Short answer")
para("We did NOT apply z-score standardization or min-max normalization to the features fed into our final XGBoost model. "
     "This is intentional and correct for the chosen algorithm.", bold=True)
spacer(3)

H2("Why no scaling for tree-based models")
para("XGBoost (our primary model) and LightGBM are tree-based ensemble methods. They make splits like:")
para('    "if lag_1 > $900,000 then go right, else go left"', italic=True)
para("The split threshold ($900,000) is found by the algorithm itself. The model is invariant to monotonic "
     "transformations — multiplying every value by 1000 or taking the log changes no decisions. So scaling "
     "adds zero value and adds computational cost.")
spacer(3)

H2("What we DID do (data preparation)")
H3("1. Type conversion")
para("All numeric columns were forced to numeric type using pd.to_numeric(errors='coerce') to handle any "
     "stray strings or formatting issues in the raw CSV.")

H3("2. Missing value handling")
para("Three different strategies were applied based on what NaN means for that column:")
bullet("Lag features: leave as NaN (drop rows where any lag is NaN — required because the first few weeks per store have no history)")
bullet("Competitor counts: fillna(0) — a missing competitor count is treated as 0 competitors")
bullet("Income brackets / housing decades: fillna(0) — a missing % is treated as 0 share")

H3("3. Categorical encoding")
para("LabelEncoder applied to Urbanicity, CBSA Type, CBSA Metro Size (see Q1 for details).")

H3("4. Sort ordering (mandatory for lag features)")
para("Before computing lags, the data is sorted by Store ID → Year → Fiscal Week. This ensures shift(n) "
     "operations correctly reach the previous N weeks of THAT store, not a different store.")
spacer(4)

H2("Where scaling WOULD matter")
para("If we ran a different algorithm family, scaling becomes necessary:")
make_table(["Algorithm", "Needs scaling?", "Why"],
    [["XGBoost (our model)", "No", "Trees split on thresholds; scale invariant"],
     ["LightGBM", "No", "Same as XGBoost"],
     ["Ridge Regression (our baseline)", "Recommended", "L2 penalty applies equally to all coefficients — features need comparable scales for fair regularization"],
     ["Linear regression (OLS)", "No (mathematically)", "Coefficients adjust to scale, but interpretation is affected"],
     ["Neural networks", "Yes (mandatory)", "Gradient descent struggles with unscaled features"],
     ["K-means / KNN", "Yes (mandatory)", "Distance metrics dominated by largest-scale feature otherwise"]],
    widths=[2.3, 1.5, 2.8])
para("For our Ridge baseline we used sklearn's default fit — coefficients are still interpretable as direction, "
     "and the WAPE comparison stays fair because all models score on the same actual sales scale.", italic=True, color=GREY)
spacer(3)

H2("Bottom line for the talking point")
para("\"Because our chosen model is tree-based (XGBoost), we deliberately did not apply z-score or min-max scaling. "
     "Trees are scale-invariant — they make split decisions on raw thresholds, so scaling would add no benefit. "
     "What we did was type-coerce, fill missing values context-appropriately, and label-encode categoricals.\"",
     italic=True)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# Q3 — HYPERPARAMETERS
# ════════════════════════════════════════════════════════════
H1("Q3 · Hyperparameters of each of the 4 models")
para("All four models trained on the same train/test split (rolling time split, 85th percentile cutoff). "
     "Hyperparameters were chosen for cross-model fairness (same learning rate + tree count for both boosting models) "
     "and to mirror industry-standard retail forecasting configurations.", italic=True, color=GREY)
spacer(4)

# Model 1
H2("1. XGBoost — Primary Production Model", color=GREEN)
make_table(["Hyperparameter", "Value", "What it does"],
    [["n_estimators", "500", "Total trees built sequentially"],
     ["learning_rate", "0.05", "Each tree's contribution to the ensemble (small = stable convergence)"],
     ["max_depth", "7", "Maximum depth of each individual tree"],
     ["subsample", "0.8", "Random row sample per tree (see Q7)"],
     ["colsample_bytree", "0.8", "Random feature sample per tree (see Q7)"],
     ["objective", "reg:squarederror", "Mean-squared-error loss for regression"],
     ["early_stopping_rounds", "50", "Stop if validation loss stops improving for 50 rounds"],
     ["random_state", "42", "Fixed seed for reproducibility"]],
    widths=[2.3, 1.5, 3.0])
spacer(3)

# Model 2
H2("2. LightGBM — Secondary Benchmark", color=BLUE)
make_table(["Hyperparameter", "Value", "What it does"],
    [["n_estimators", "500", "Same as XGBoost for fair comparison"],
     ["learning_rate", "0.05", "Same as XGBoost"],
     ["max_depth", "8", "Slightly deeper because LightGBM grows leaf-wise"],
     ["num_leaves", "63", "Maximum leaves per tree (key LightGBM-specific control)"],
     ["min_child_samples", "50", "Minimum samples required in a leaf — prevents tiny over-fit splits"],
     ["subsample", "0.8", "Same as XGBoost"],
     ["colsample_bytree", "0.8", "Same as XGBoost"],
     ["objective", "regression", "Standard regression objective"],
     ["early_stopping_rounds", "50", "Same as XGBoost"],
     ["random_state", "42", "Same seed for reproducibility"]],
    widths=[2.3, 1.5, 3.0])
spacer(3)

# Model 3
H2("3. Ridge Regression — Linear Baseline", color=PURPLE)
make_table(["Hyperparameter", "Value", "What it does"],
    [["alpha", "1.0", "L2 regularization strength (default sklearn value)"],
     ["fit_intercept", "True", "Allow non-zero intercept"],
     ["solver", "auto", "sklearn chooses based on data shape"],
     ["random_state", "42", "Reproducibility"]],
    widths=[2.3, 1.5, 3.0])
para("Ridge serves as our linear sanity check. If the gradient-boosting models couldn't beat regularized "
     "linear regression, the added complexity wouldn't be justified. Ridge scored WAPE 7.11% — very close to XGBoost's 7.05%, "
     "telling us the relationship is fairly linear once lag features are present, but the non-linear models still edge it out.",
     italic=True, color=GREY)
spacer(3)

# Model 4
H2("4. Naive Lag-52 Baseline — Zero Hyperparameters", color=ORANGE)
para("The simplest possible 'model': prediction for week t = actual sales at week t-52 (same week last year). "
     "There are no parameters to tune. This is the bar every real model must clear to justify its existence. "
     "WAPE = 8.27%. Our final model beats it by 15%.")
spacer(4)

H2("Head-to-head holdout results")
make_table(["Model", "WAPE", "Bias", "R²", "Verdict"],
    [["XGBoost (selected)", "7.05%", "−0.97%", "0.9769", "Best balanced"],
     ["LightGBM", "7.07%", "−1.11%", "0.9773", "Confirms robustness"],
     ["Ridge", "7.11%", "+0.39%", "0.9783", "Strong linear baseline"],
     ["Naive (lag-52)", "8.27%", "+0.33%", "0.9648", "The bar to beat"]],
    widths=[2.0, 1.0, 1.0, 1.0, 2.0])

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# Q4 — WAPE + FEATURE WEIGHTS
# ════════════════════════════════════════════════════════════
H1("Q4 · Why WAPE was chosen + the final feature weights")
spacer(4)

H2("Part A · Why WAPE (not MAPE, RMSE, or MAE)")
para("WAPE = Weighted Absolute Percentage Error. Formula:", bold=True)
para("    WAPE = Σ |Actual − Predicted| / Σ Actual × 100", italic=True)

H3("Plain-English meaning")
para("Of every $100 of actual sales, how many dollars did the model miss by? Our 7.05% → about $7 of error "
     "per $100 sold.")
spacer(3)

H3("Why WAPE beats the alternatives for retail")
make_table(["Metric", "Issue", "Why WAPE wins"],
    [["MAPE", "Equal weight to a $1K store and a $50M store; explodes when actual is near zero",
              "WAPE weights by sales volume — bigger stores' errors count more, matching dollar risk"],
     ["RMSE", "Penalizes large errors quadratically; hard to interpret in dollar terms",
              "WAPE is a single intuitive percentage"],
     ["MAE (dollars)", "No sense of scale — is $50K error big or small?",
              "WAPE normalizes to % of actual, immediately scale-aware"],
     ["Bias %", "Only direction, not magnitude",
              "WAPE measures magnitude — we report bias separately as a second axis"]],
    widths=[1.0, 3.0, 3.0])
spacer(3)

H3("Industry standard")
para("WAPE is the de-facto reporting metric for retail demand forecasting (used by Walmart, Amazon, Lowe's "
     "own legacy planning) precisely because it preserves dollar weighting and yields a single intuitive number.")
spacer(6)

# Part B — Feature importance
H2("Part B · The 16 feature weights (XGBoost importance)")
para("These are the relative importance scores XGBoost assigned each feature based on how often and how "
     "decisively the feature was used to split data across all 500 trees. Importance is normalized to sum to 100%.",
     italic=True, color=GREY)
spacer(3)

make_table(["Rank", "Feature", "Group", "Importance %"],
    [["1", "lag_1", "Lag / Momentum", "62.22%"],
     ["2", "lag_4", "Lag / Momentum", "29.63%"],
     ["3", "lag_52", "Lag / Momentum", "3.14%"],
     ["4", "Wallflowers Depot (TA)", "Competition", "0.70%"],
     ["5", "lag_13", "Lag / Momentum", "0.55%"],
     ["6", "housing_new_share", "Demand – Engineered", "0.40%"],
     ["7", "CYE Total Population", "Demand – Raw", "0.39%"],
     ["8", "CYE Median HH Income", "Demand – Raw", "0.38%"],
     ["9", "CYE Total Housing Units", "Demand – Raw", "0.36%"],
     ["10", "income_affluent", "Demand – Engineered", "0.35%"],
     ["11", "housing_old_share", "Demand – Engineered", "0.34%"],
     ["12", "CYE HH Density / SqMi", "Demand – Raw", "0.34%"],
     ["13", "Sister Store Count (TA)", "Competition", "0.34%"],
     ["14", "total_competitor_ta", "Competition", "0.33%"],
     ["15", "CYE Total Households", "Demand – Raw", "0.29%"],
     ["16", "Urbanicity (encoded)", "Store / Market", "0.24%"]],
    widths=[0.6, 2.3, 1.9, 1.2])
spacer(3)

H3("How to read the distribution")
bullet("Lag features dominate (95.5% combined): lag_1 + lag_4 = 91.85%. This is expected and correct for one-week-ahead forecasting — recent sales are the strongest predictor of next-week sales.")
bullet("Market features (housing, income, competition, population) collectively carry ~4% — small but non-zero. They differentiate stores when the lag signal is ambiguous.")
bullet("This distribution is healthy: a planner can defend any target by pointing to 'recent run rate, with adjustments for the local market context'.")

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# Q5 — 5 STORE GROUPS
# ════════════════════════════════════════════════════════════
H1("Q5 · How the 5 store-role groups are chosen + thresholds")
spacer(4)

H2("The 5 segments and what they mean")
make_table(["Role", "What it signals", "Planner playbook"],
    [["High Growth", "Strong demand + store outperforming model", "Raise targets aggressively"],
     ["Growth", "Moderately positive market signals", "Use data-driven target as-is"],
     ["Neutral", "Balanced — no strong signal either way", "Hold pattern; standard target"],
     ["Maintain", "Stable demand, moderate competition", "Hold targets, monitor closely"],
     ["Defend", "High competition + store underperforming", "Conservative targets; protect share"]],
    widths=[1.3, 2.5, 2.6])
spacer(3)

H2("The two composite scores")
para("Each store gets two derived scores from its 16 features:", bold=True)
para('    growth_score = (CAGR_HH / 2) + (housing_new_share / 10) − (Bias_pct / 10)', italic=True)
para('    risk_score   = (total_competitor_ta / 50) + (Bias_pct / 10)', italic=True)
para('    net_score    = growth_score − risk_score', italic=True)
spacer(2)

H3("What each component captures")
bullet("CAGR HH 2010-2020 (positive) → household growth = forward demand expansion")
bullet("housing_new_share (positive) → new construction → new-homeowner improvement demand")
bullet("Negative bias % → model UNDER-predicts → store is outperforming → upward signal")
bullet("total_competitor_ta (positive) → more competition = more downward pressure")
bullet("Positive bias % → model OVER-predicts → store is underperforming → downward signal")
spacer(3)

H2("Why QUANTILE-based thresholds, not fixed cut-offs")
para("Our first iteration used fixed thresholds (e.g., growth_score > 1.5 = High Growth). Result was useless: "
     "86% of stores landed in 'High Growth'. The growth_score distribution was right-skewed and most stores "
     "naturally scored above the fixed line.")
para("Switched to quantile-based binning, which guarantees ~20% of stores in each segment regardless of "
     "underlying distribution shape.")
spacer(3)

H2("The thresholds (quantile-based)")
make_table(["Segment", "Rule", "Stores", "% of total"],
    [["High Growth", "net_score ≥ 80th percentile", "346", "20.0%"],
     ["Growth", "60th to 80th percentile", "345", "20.0%"],
     ["Neutral", "40th to 60th percentile", "345", "20.0%"],
     ["Maintain", "20th to 40th percentile", "346", "20.0%"],
     ["Defend", "Below 20th percentile", "345", "20.0%"]],
    widths=[1.5, 2.5, 1.2, 1.3])
para("Total: 1,727 stores classified, ~20% in each of 5 segments.", italic=True, color=GREY)
spacer(4)

H2("Critical caveat — this is a heuristic, not a learned classifier")
para("Role-of-Store is a rule-based strategic OVERLAY, not a separate trained ML model. The composite scores "
     "use hand-chosen weights (divide-by-2 for CAGR, divide-by-10 for housing share, etc.) tuned for "
     "interpretability rather than optimality. A future iteration could replace this with a learned clustering "
     "algorithm (K-Means on standardized features).", italic=True)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# Q6 — MULTICOLLINEARITY DROPS
# ════════════════════════════════════════════════════════════
H1("Q6 · Features ignored due to multicollinearity")
spacer(4)

H2("Categories of dropped features")

H3("Category A · Sum-to-100% groups (perfect multicollinearity)")
para("These columns within a group sum to ~100 for every store, so they're mathematically redundant.")
spacer(2)

para("Income brackets — 8 columns reduced to 3 strategic shares + median:", bold=True)
make_table(["Dropped column", "Why redundant"],
    [["CYE % HH Income 25k-35k", "Captured by income_low_share aggregate"],
     ["CYE % HH Income 35k-50k", "Mid-low bracket — collapsed into low/diy aggregates"],
     ["CYE % HH Income 50k-75k", "Captured by income_diy_core (dropped in iteration 3)"],
     ["CYE % HH Income 75k-100K", "Captured by income_diy_core (dropped in iteration 3)"],
     ["CYE % HH Income 100k-150k", "Middle income — captured by median + adjacent aggregates"],
     ["CYE % HH Income 150k-250k", "Captured by income_affluent aggregate"],
     ["CYE % HH Income 250K Plus", "Captured by income_affluent aggregate (also 41.6% missing)"]],
    widths=[3.0, 3.8])
spacer(3)

para("Housing decade buckets — 9 columns reduced to median + 2 strategic shares:", bold=True)
make_table(["Dropped column", "Why redundant"],
    [["CYE % HU Built 1970-1979", "Middle decade — captured by Median Year Built"],
     ["CYE % HU Built 1980-1989", "Middle decade — captured by Median Year Built"],
     ["CYE % HU Built 1990-1999", "Middle decade — captured by Median Year Built"],
     ["CYE % HU Built 2020+", "Very small share in most areas — sparse signal"]],
    widths=[3.0, 3.8])
spacer(3)

H3("Category B · Perfect inverses")
para("CYE % Housing Units Rented = 100% − CYE % Housing Units Owned. Perfect linear redundancy → one is "
     "dropped (we kept % Owned because ownership drives improvement spend).")
spacer(3)

H3("Category C · Aggregate vs component redundancy")
make_table(["Dropped", "Why"],
    [["CYE Aggregate Family Household Income", "Equals Median × Total Households (redundant)"],
     ["CYE Average Mean Household Income", "Mean is sensitive to outliers; Median is more robust"]],
    widths=[3.2, 3.6])
spacer(3)

H3("Category D · Distance-band duplicates (~80 columns dropped)")
para("Each competitor has 5-6 distance band columns (Trade Area, 3mi, 5mi, 10mi, 0-11mi, 11-30mi) and they "
     "are heavily correlated. We kept TradeArea only — per Lowe's own methodology, the proprietary Trade Area "
     "captures ~70% of a store's true sales catchment, making it the gold standard.")
spacer(3)

H3("Category E · Rolling features dropped in iteration 3")
para("roll_4 and roll_13 were initially in the 40-feature set. After the first model iteration they accounted "
     "for ~44% of feature importance combined, masking the contributions of market features. We dropped them in "
     "iteration 3, accepting a small WAPE increase (6.21 → 7.05%) in exchange for:", italic=True)
bullet("Bias improvement from −2.53% to −0.97% (near-zero)")
bullet("Better explainability — the model's drivers are now interpretable to a planner")
bullet("More balanced feature importance distribution")
spacer(3)

H2("Why these were 'less significant'")
para("In all cases the dropped features carried duplicate or near-duplicate information. Key reasons:", bold=True)
numbered("In linear models (Ridge), multicollinearity destabilizes coefficient estimates and inflates standard errors")
numbered("In tree models, correlated features compete for splits — one gets used, others sit idle, but they each add training cost and complicate feature-importance interpretation")
numbered("Fewer features means simpler explanations to business stakeholders (planners must defend targets)")
numbered("Reduces overfitting risk by lowering the model's effective degrees of freedom")
numbered("Faster training and inference")
spacer(3)

H2("The funnel summary")
para("From 165 candidate features (after identity & leakage drops) to the final 16 — multicollinearity removal "
     "alone accounted for ~100 of those dropped features (distance bands + income brackets + housing decades + "
     "inverses + redundant aggregates).")

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# Q7 — subsample / colsample explained
# ════════════════════════════════════════════════════════════
H1("Q7 · 'subsample 0.8' and 'colsample_bytree 0.8' in plain English")
spacer(4)

H2("The intuition: building a panel of judges")
para("Imagine you're building a panel of 500 judges to make a decision (each judge = one tree). If all 500 "
     "judges had the exact same information and saw all the same evidence, they would make highly correlated "
     "decisions — and the panel's overall judgment wouldn't be much better than one expert. ", italic=True)
para("XGBoost's solution: deliberately introduce randomness so each judge gets a slightly different view of the "
     "data, producing diverse opinions that average into a more robust final answer.", italic=True)
spacer(4)

H2("subsample = 0.8")
H3("What it does")
para("When building each individual tree, randomly select 80% of the training rows (store-week observations) "
     "and use ONLY those rows to fit that tree. The other 20% are temporarily ignored for that tree only — they "
     "may be used by the next tree.")
spacer(2)

H3("Layman analogy")
para("Think of it like getting feedback from 100 customers, but each focus group only invites a random 80 "
     "of them. Each focus group hears slightly different stories. The combined output across many focus groups "
     "is more representative than any single full-group session that one strong voice might dominate.")
spacer(2)

H3("Why 0.8 specifically (not 1.0, not 0.5)")
bullet("1.0 means every tree sees all data → trees become too similar, ensemble loses diversity benefit, overfits to training set")
bullet("0.5 means each tree sees only half — too aggressive, each tree becomes weaker, noisier")
bullet("0.8 is the industry-standard sweet spot — enough diversity to regularize, enough data per tree to learn well")
spacer(2)

H3("What it prevents")
para("Overfitting. Without subsampling, an ensemble of 500 trees can memorize quirks of the exact training data "
     "(e.g., one particular store-week that had an unusual sale). Random subsampling makes the model robust to "
     "such quirks.")
spacer(5)

H2("colsample_bytree = 0.8")
H3("What it does")
para("When building each individual tree, randomly select 80% of the FEATURES and consider only those features "
     "for the splits in that tree. The other 20% of features are unavailable to that tree.")
spacer(2)

H3("Layman analogy")
para("Imagine each judge in your panel is given only 80% of the available evidence to make their decision. "
     "Judge A might be shown lag_1, income, and competition but NOT housing share. Judge B might be shown "
     "housing share, urbanicity, and lag_4 but NOT lag_1. By forcing different judges to consider different "
     "evidence, you prevent any single piece of evidence from dominating every decision.")
spacer(2)

H3("Why this matters HUGELY for our model specifically")
para("Our lag_1 feature is so dominant (correlation 0.982 with target) that without colsample_bytree, every "
     "single one of the 500 trees would split on lag_1 first. Other features would never get a chance to "
     "contribute.")
para("With colsample_bytree=0.8, lag_1 is unavailable in ~20% of trees. Those trees are FORCED to find their "
     "second-best feature — and that's exactly when market signals (housing, income, competition) get to "
     "shine. The end result: a model whose feature importance distribution is meaningful, not monopolized.", italic=True)
spacer(3)

H3("Effect on our reported feature importances")
para("Compare what you'd see WITH vs WITHOUT colsample_bytree=0.8 (illustrative):")
make_table(["Feature", "With colsample 1.0 (hypothetical)", "With colsample 0.8 (our model)"],
    [["lag_1", "~92%", "62.22%"],
     ["lag_4", "~5%", "29.63%"],
     ["Market features combined", "~3%", "~5%"],
     ["Explainability", "Black box (one feature drives everything)", "Defensible (4% explained by interpretable market signals)"]],
    widths=[2.2, 2.3, 2.3])
spacer(4)

H2("Combined effect — why both matter")
para("Both parameters introduce randomness, but along different axes:", bold=True)
bullet("subsample randomizes WHICH OBSERVATIONS each tree learns from (rows)")
bullet("colsample_bytree randomizes WHICH FEATURES each tree can use (columns)")
para("Together they make the 500-tree ensemble truly diverse — and reduced overfitting is the result. "
     "Both are forms of 'regularization through randomness'.", italic=True)
spacer(4)

H2("Bottom line for the talking point")
para("\"subsample 0.8 means each tree only sees 80% of the rows, randomly chosen. colsample_bytree 0.8 means "
     "each tree only considers 80% of the features. Both inject randomness so the 500 trees in our ensemble "
     "are diverse instead of identical. The end result is a model that generalizes well to new data and doesn't "
     "let any single feature dominate every decision.\"", italic=True)

# Footer
spacer(8)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Capstone 8 · Group 8 · IIM Calcutta APAL02   |   github.com/nishthaspeakx/capstone_8")
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY

# Save & convert
docx_path = "outputs/Capstone8_Geetha_QA.docx"
doc.save(docx_path)
print(f"Saved DOCX: {docx_path}")
result = subprocess.run(
    ["soffice","--headless","--convert-to","pdf",docx_path,"--outdir","outputs"],
    capture_output=True, text=True, timeout=180)
print((result.stdout or result.stderr)[-300:])
pdf_path = docx_path.replace(".docx",".pdf")
desktop = os.path.expanduser("~/Desktop/Capstone8_Geetha_QA.pdf")
shutil.copy(pdf_path, desktop)
print(f"PDF: {pdf_path}")
print(f"Desktop: {desktop}")
print(f"Size: {os.path.getsize(pdf_path)/1024:.1f} KB")
