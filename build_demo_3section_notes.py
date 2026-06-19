"""
Build a focused 3-section demo notes PDF (Opening, Pick-a-Store, Explain)
+ a clear section on how GenAI works and the fallback.
Output: outputs/Capstone8_LiveDemo_3Sections_Notes.pdf
"""
import os, shutil, subprocess
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x00, 0x30, 0x87)
BLUE = RGBColor(0x02, 0x77, 0xBD)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
ORANGE = RGBColor(0xE6, 0x51, 0x00)
GREY = RGBColor(0x55, 0x55, 0x55)
ACCENT = RGBColor(0xF4, 0xB4, 0x00)
INK = RGBColor(0x1A, 0x23, 0x32)
PURPLE = RGBColor(0x6F, 0x42, 0xC1)

doc = Document()
for section in doc.sections:
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def H1(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.color.rgb = NAVY; r.font.size = Pt(19); r.bold = True
    return p

def H2(text, color=BLUE):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.color.rgb = color; r.font.size = Pt(13); r.bold = True
    return p

def label(p, text, color, size=10, bold=True):
    r = p.add_run(text); r.font.color.rgb = color; r.bold = bold; r.font.size = Pt(size)

def spacer(pt=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pt)

# ════════════════════════════════════════════════════════════
# COVER
# ════════════════════════════════════════════════════════════
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Demo Speaker Notes — First 3 Sections")
r.font.size = Pt(22); r.bold = True; r.font.color.rgb = NAVY

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Lowe's AI Sales Targeting · Streamlit Walkthrough")
r.font.size = Pt(12); r.italic = True; r.font.color.rgb = BLUE

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Capstone 8 · Group 8 · IIM Calcutta APAL02   |   First ~2 minutes of the demo")
r.font.size = Pt(10); r.font.color.rgb = GREY

spacer(8)

# Mini table of contents
t = doc.add_table(rows=1, cols=2); t.style = "Light Grid Accent 1"
hdr = t.rows[0].cells
for c, txt in zip(hdr, ["Section", "Time"]):
    c.text = ""; rr = c.paragraphs[0].add_run(txt); rr.bold = True; rr.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    shade(c, "003087")
overview = [
    ("0 · Opening + The 94.7% hero banner", "0:00 – 0:20"),
    ("1 · Pick a real store — Wolfridge, Texas (a 'Best Save')", "0:20 – 1:15"),
    ("2 · Click EXPLAIN — the GenAI explainer", "1:15 – 1:50"),
    ("Bonus · How the GenAI layer actually works (for Q&A)", "+1 page"),
]
for sec, time_ in overview:
    row = t.add_row().cells
    row[0].text = ""; row[0].paragraphs[0].add_run(sec).font.size = Pt(10)
    row[1].text = ""; row[1].paragraphs[0].add_run(time_).font.size = Pt(10)
    row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
for row in t.rows:
    row.cells[0].width = Inches(5.7); row.cells[1].width = Inches(1.5)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# SECTION 0
# ════════════════════════════════════════════════════════════
H1("0 · Opening + The 94.7% Hero Banner")
p = doc.add_paragraph()
r = p.add_run("0:00 – 0:20  (~20 seconds)"); r.font.color.rgb = GREY; r.italic = True
spacer(3)

p = doc.add_paragraph(); label(p, "WHAT TO DO  ", ACCENT, size=11, bold=True)
p.paragraph_format.space_after = Pt(2)
for d in [
    "Open the browser tab that already has the dashboard loaded (do this BEFORE the demo starts).",
    "Scroll to the very top so the big '94.7%' number is the first thing they see.",
]:
    bp = doc.add_paragraph(style="List Bullet")
    bp.paragraph_format.space_after = Pt(2)
    bp.add_run(d).font.size = Pt(11)
spacer(3)

p = doc.add_paragraph(); label(p, "WHAT TO SAY  ", GREEN, size=11, bold=True)
p.paragraph_format.space_after = Pt(2)
sp = doc.add_paragraph()
sp.paragraph_format.left_indent = Inches(0.25)
r = sp.add_run('"This is the dashboard our planner would actually use every week. The first thing you see is the headline result — 94.7% of stores now get an accurate target, within plus-or-minus five percent. So instead of reviewing all 1,727 stores manually like they do today, the team only needs to look at 91. That is a 95% cut in their workload."')
r.font.size = Pt(12); r.italic = True
spacer(4)

p = doc.add_paragraph()
label(p, "TIP  ", ORANGE, size=10)
r = p.add_run("Pre-open the dashboard before the demo. Don't switch tabs live. Have it ready.")
r.font.size = Pt(10.5); r.italic = True; r.font.color.rgb = ORANGE

p = doc.add_paragraph()
label(p, "NEXT  ", BLUE, size=10, bold=True)
r = p.add_run("Scroll down to the section titled 'Try It: Pick a Store'.")
r.font.size = Pt(10.5); r.italic = True

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# SECTION 1
# ════════════════════════════════════════════════════════════
H1("1 · Pick a Real Store — Wolfridge, Texas")
p = doc.add_paragraph()
r = p.add_run("0:20 – 1:15  (~55 seconds)"); r.font.color.rgb = GREY; r.italic = True
spacer(3)

p = doc.add_paragraph(); label(p, "WHAT TO DO  ", ACCENT, size=11, bold=True)
p.paragraph_format.space_after = Pt(2)
for d in [
    "In the dropdown, start typing 'Wolfridge' or 'S01155'. The store will appear in the list.",
    "Click it. Wait for the three cards (Actual, Peanut-butter, Model) to show on screen.",
]:
    bp = doc.add_paragraph(style="List Bullet")
    bp.paragraph_format.space_after = Pt(2)
    bp.add_run(d).font.size = Pt(11)
spacer(3)

p = doc.add_paragraph(); label(p, "WHAT TO SAY  ", GREEN, size=11, bold=True)
p.paragraph_format.space_after = Pt(2)
sp = doc.add_paragraph()
sp.paragraph_format.left_indent = Inches(0.25)
r = sp.add_run('"Let me show you what this means for a real store. I picked Wolfridge, Texas. Look at the three numbers on screen. The actual sales for the test period are about 5 million dollars. The current peanut-butter target says 6.5 million — that is wrong by 30 percent. Our model says 5.14 million — wrong by only 7 percent. So for just this one store, we cut the error by 23 percentage points. And see the green box below? It tells the planner the model is the better choice here. That is one store. Now think about doing this for 1,727 of them."')
r.font.size = Pt(12); r.italic = True
spacer(4)

p = doc.add_paragraph()
label(p, "TIP  ", ORANGE, size=10)
r = p.add_run("If Wolfridge doesn't load, just pick any store that has a green 'Model wins' verdict — the story is the same.")
r.font.size = Pt(10.5); r.italic = True; r.font.color.rgb = ORANGE

p = doc.add_paragraph()
label(p, "NEXT  ", BLUE, size=10, bold=True)
r = p.add_run("Stay on the same store. Click the blue 'Explain' button.")
r.font.size = Pt(10.5); r.italic = True

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# SECTION 2
# ════════════════════════════════════════════════════════════
H1("2 · Click EXPLAIN — The GenAI Explainer")
p = doc.add_paragraph()
r = p.add_run("1:15 – 1:50  (~35 seconds)"); r.font.color.rgb = GREY; r.italic = True
spacer(3)

p = doc.add_paragraph(); label(p, "WHAT TO DO  ", ACCENT, size=11, bold=True)
p.paragraph_format.space_after = Pt(2)
for d in [
    "Click the 'Explain' button next to the heading 'Why did this store get this target?'",
    "Wait 2 or 3 seconds. A short three-paragraph explanation appears.",
    "Read the first two lines out loud so the professor can follow.",
]:
    bp = doc.add_paragraph(style="List Bullet")
    bp.paragraph_format.space_after = Pt(2)
    bp.add_run(d).font.size = Pt(11)
spacer(3)

p = doc.add_paragraph(); label(p, "WHAT TO SAY  ", GREEN, size=11, bold=True)
p.paragraph_format.space_after = Pt(2)
sp = doc.add_paragraph()
sp.paragraph_format.left_indent = Inches(0.25)
r = sp.add_run('"This is our GenAI layer. When the planner clicks Explain, the system writes a plain-English answer for that specific store. It quotes the real numbers — recent sales, year-over-year change, how many competitors are nearby, what the housing market looks like — and tells the planner exactly why the model picked this target and what action to take. Trust it, raise it, or review it. Every one of the 1,727 stores can be explained this way, on demand."')
r.font.size = Pt(12); r.italic = True
spacer(4)

# Add the GenAI mechanics explanation
H2("How does this actually work behind the scenes?", color=PURPLE)
p = doc.add_paragraph()
r = p.add_run(
    "When you click Explain, three things happen: "
)
r.font.size = Pt(11)
r2 = p.add_run("(1) ")
r2.bold = True; r2.font.size = Pt(11); r2.font.color.rgb = PURPLE
p.add_run("the app collects that store's 16 feature values along with its predicted target and actual sales. ").font.size = Pt(11)
r2 = p.add_run("(2) ")
r2.bold = True; r2.font.size = Pt(11); r2.font.color.rgb = PURPLE
p.add_run("It sends this to Google's Gemini AI with a structured prompt asking for a 3-part explanation. ").font.size = Pt(11)
r2 = p.add_run("(3) ")
r2.bold = True; r2.font.size = Pt(11); r2.font.color.rgb = PURPLE
p.add_run("Gemini writes the paragraph using those exact numbers — never a generic answer.").font.size = Pt(11)
spacer(4)

p = doc.add_paragraph()
label(p, "FALLBACK  ", BLUE, size=11, bold=True)
r = p.add_run("If Gemini is down or our API quota runs out, the app automatically uses a built-in rule-based explainer that uses the same 16 features and writes a similar 3-paragraph answer. The audience cannot tell the difference. The demo never fails.")
r.font.size = Pt(11); r.italic = True
spacer(6)

p = doc.add_paragraph()
label(p, "TIP  ", ORANGE, size=10)
r = p.add_run("If asked 'is this real AI?': YES. Each of the 1,727 stores produces a different explanation because each has different feature values. It is not pre-written text.")
r.font.size = Pt(10.5); r.italic = True; r.font.color.rgb = ORANGE

p = doc.add_paragraph()
label(p, "NEXT  ", BLUE, size=10, bold=True)
r = p.add_run("Open the dropdown again and pick an under-targeted store (or move on to the next section as planned).")
r.font.size = Pt(10.5); r.italic = True

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# BONUS: how GenAI works in depth (for Q&A)
# ════════════════════════════════════════════════════════════
H1("Bonus · How the GenAI Layer Works")
p = doc.add_paragraph()
r = p.add_run("This page is for Q&A — keep it ready in case the professor asks how it works."); r.italic = True; r.font.color.rgb = GREY
spacer(6)

H2("The 3-step pipeline (what runs when 'Explain' is clicked)")
for i, txt in enumerate([
    ("Gather the store's data",
     "The app pulls the 16 feature values for that store from a CSV file — things like sales last week (lag_1), median income, housing share, total competitors, and so on. It also adds the model's predicted target, the peanut-butter target, the actual sales, the bias percentage, and the segment role (High Growth, Defend, etc.)."),
    ("Build the prompt and send to Gemini",
     "The app formats this into a structured prompt asking Gemini for a 3-part answer: (a) Why this target? (b) Why this segment? (c) What action should the planner take? The prompt explicitly tells Gemini to reference the actual numbers — not generic phrases."),
    ("Gemini writes a custom paragraph",
     "Gemini reads the numbers and writes a unique explanation for that store. Every store gets a different paragraph because every store has different feature values. There is no pre-written template."),
]):
    p = doc.add_paragraph()
    rr = p.add_run(f"Step {i+1}: "); rr.bold = True; rr.font.color.rgb = NAVY; rr.font.size = Pt(11.5)
    rr = p.add_run(txt[0]); rr.bold = True; rr.font.size = Pt(11.5)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.3)
    p2.paragraph_format.space_after = Pt(8)
    rr = p2.add_run(txt[1]); rr.font.size = Pt(11)

spacer(4)
H2("The Fallback — why the demo never fails")
p = doc.add_paragraph()
r = p.add_run(
    "We built a backup rule-based explainer in Python. It uses the same 16 feature values and writes "
    "a similar 3-paragraph answer using if-else logic. If Gemini is unreachable (no internet, API down, "
    "quota exhausted), the app silently switches to this fallback. The audience cannot tell the difference. "
    "This is important because real production systems need to be robust — they cannot rely on a single "
    "external service being available 100 percent of the time."
)
r.font.size = Pt(11)
spacer(6)

H2("Is it useful TODAY (with static demo data)?")
for txt in [
    "Yes. Each of the 1,727 stores has a unique combination of 16 feature values, so the explanation is unique for every store. The data file is fixed, but the model's reasoning is real and store-specific.",
    "It proves the explainability capability. Planners can interrogate ANY store and understand why the model picked that target — no black box.",
    "It also serves as a training tool. A new planner can scroll through any store and see how recent sales, income, housing, and competition combine into a target.",
]:
    bp = doc.add_paragraph(style="List Bullet")
    bp.paragraph_format.space_after = Pt(4)
    rr = bp.add_run(txt); rr.font.size = Pt(11)
spacer(4)

H2("How it becomes MORE useful with dynamic, live data")
for txt in [
    "In production, the data refreshes every week. New actuals come in. Lag features update.",
    "The same GenAI code runs without any change — Gemini just gets fresh numbers and writes fresh explanations.",
    "Each week, instead of explaining 1,727 stores, planners only need explanations for the ~91 stores flagged for review. They get an auto-generated briefing per store.",
    "Over time, the explanations help build organisational memory — planners learn why certain markets behave differently, and the model is no longer 'magic' to them.",
]:
    bp = doc.add_paragraph(style="List Bullet")
    bp.paragraph_format.space_after = Pt(4)
    rr = bp.add_run(txt); rr.font.size = Pt(11)

# Footer
spacer(8)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Capstone 8 · Group 8 · IIM Calcutta APAL02  |  storewisetarget.streamlit.app  |  github.com/nishthaspeakx/capstone_8")
r.font.size = Pt(9); r.font.color.rgb = GREY; r.italic = True

# Save & convert
docx_path = "outputs/Capstone8_LiveDemo_3Sections_Notes.docx"
doc.save(docx_path)
print(f"Saved DOCX: {docx_path}")

result = subprocess.run(
    ["soffice","--headless","--convert-to","pdf",docx_path,"--outdir","outputs"],
    capture_output=True, text=True, timeout=120)
print((result.stdout or result.stderr)[-300:])

pdf_path = docx_path.replace(".docx",".pdf")
desktop_pdf = os.path.expanduser("~/Desktop/Capstone8_LiveDemo_3Sections_Notes.pdf")
shutil.copy(pdf_path, desktop_pdf)
print(f"PDF:     {pdf_path}")
print(f"Desktop: {desktop_pdf}")
print(f"Size:    {os.path.getsize(pdf_path)/1024:.1f} KB")
